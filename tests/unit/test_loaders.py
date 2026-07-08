import io
from unittest.mock import MagicMock, patch

import pytest
from docx import Document as DocxDocument

from rag.exceptions import LoaderError
from rag.ingestion.loaders import (
    DocxLoader,
    HtmlLoader,
    LoaderFactory,
    MarkdownLoader,
    PdfLoader,
    TextLoader,
    normalize_text,
)
from rag.models import DocumentFormat


def test_normalize_text_collapses_whitespace() -> None:
    raw = "a   b\t c\n\n\n\n\nd  \n"
    assert normalize_text(raw) == "a b c\n\nd"


def test_text_loader() -> None:
    assert TextLoader().load(b"hello  world") == "hello world"


def test_text_loader_rejects_empty() -> None:
    with pytest.raises(LoaderError, match="no text content"):
        TextLoader().load(b"   \n  ")


def test_html_loader_strips_scripts() -> None:
    html = b"<html><body><script>evil()</script><p>visible text</p></body></html>"
    result = HtmlLoader().load(html)
    assert "visible text" in result
    assert "evil" not in result


def test_markdown_loader_strips_formatting() -> None:
    md = b"# Title\n\nSome **bold** and `code`.\n\n- item one\n- item two"
    result = MarkdownLoader().load(md)
    assert "Title" in result
    assert "**" not in result
    assert "item one" in result


def test_docx_loader_paragraphs_and_tables() -> None:
    document = DocxDocument()
    document.add_paragraph("First paragraph")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "cell a"
    table.rows[0].cells[1].text = "cell b"
    buffer = io.BytesIO()
    document.save(buffer)
    result = DocxLoader().load(buffer.getvalue())
    assert "First paragraph" in result
    assert "cell a" in result
    assert "cell b" in result


def test_pdf_loader_joins_pages() -> None:
    page_one = MagicMock()
    page_one.extract_text.return_value = "page one text"
    page_two = MagicMock()
    page_two.extract_text.return_value = None
    reader = MagicMock()
    reader.pages = [page_one, page_two]
    with patch("rag.ingestion.loaders.PdfReader", return_value=reader):
        result = PdfLoader().load(b"%PDF-fake")
    assert result == "page one text"


def test_pdf_loader_wraps_parse_errors() -> None:
    with pytest.raises(LoaderError, match="failed to parse pdf"):
        PdfLoader().load(b"not a pdf at all")


@pytest.mark.parametrize(
    ("filename", "expected"),
    [
        ("report.PDF", DocumentFormat.PDF),
        ("notes.md", DocumentFormat.MARKDOWN),
        ("page.htm", DocumentFormat.HTML),
        ("doc.docx", DocumentFormat.DOCX),
        ("plain.txt", DocumentFormat.TEXT),
        ("dir/sub/file.markdown", DocumentFormat.MARKDOWN),
    ],
)
def test_detect_format(filename: str, expected: DocumentFormat) -> None:
    assert LoaderFactory.detect_format(filename) is expected


def test_detect_format_unsupported() -> None:
    with pytest.raises(LoaderError, match="unsupported file extension"):
        LoaderFactory.detect_format("archive.zip")


def test_factory_returns_matching_loader() -> None:
    for format in DocumentFormat:
        assert LoaderFactory.for_format(format).format is format
