import io
import re
from abc import ABC, abstractmethod
from pathlib import PurePosixPath, PureWindowsPath
from typing import ClassVar

from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from markdown_it import MarkdownIt
from pypdf import PdfReader

from rag.exceptions import LoaderError
from rag.models import DocumentFormat

_BLANK_LINES = re.compile(r"\n{3,}")


def normalize_text(text: str) -> str:
    lines = [" ".join(line.split()) for line in text.splitlines()]
    collapsed = _BLANK_LINES.sub("\n\n", "\n".join(lines))
    return collapsed.strip()


class DocumentLoader(ABC):
    format: ClassVar[DocumentFormat]

    @abstractmethod
    def extract_text(self, data: bytes) -> str: ...

    def load(self, data: bytes) -> str:
        try:
            text = self.extract_text(data)
        except LoaderError:
            raise
        except Exception as exc:
            raise LoaderError(f"failed to parse {self.format.value} document") from exc
        cleaned = normalize_text(text)
        if not cleaned:
            raise LoaderError(f"no text content in {self.format.value} document")
        return cleaned


class PdfLoader(DocumentLoader):
    format = DocumentFormat.PDF

    def extract_text(self, data: bytes) -> str:
        reader = PdfReader(io.BytesIO(data))
        return "\n\n".join(page.extract_text() or "" for page in reader.pages)


class DocxLoader(DocumentLoader):
    format = DocumentFormat.DOCX

    def extract_text(self, data: bytes) -> str:
        document = DocxDocument(io.BytesIO(data))
        parts = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            parts.extend("\t".join(cell.text for cell in row.cells) for row in table.rows)
        return "\n".join(parts)


class HtmlLoader(DocumentLoader):
    format = DocumentFormat.HTML

    def extract_text(self, data: bytes) -> str:
        soup = BeautifulSoup(data, "html.parser")
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()
        return soup.get_text(separator="\n")


class MarkdownLoader(DocumentLoader):
    format = DocumentFormat.MARKDOWN

    def extract_text(self, data: bytes) -> str:
        rendered = MarkdownIt().render(data.decode("utf-8"))
        soup = BeautifulSoup(rendered, "html.parser")
        return soup.get_text(separator="\n")


class TextLoader(DocumentLoader):
    format = DocumentFormat.TEXT

    def extract_text(self, data: bytes) -> str:
        return data.decode("utf-8")


_SUFFIX_FORMATS: dict[str, DocumentFormat] = {
    ".pdf": DocumentFormat.PDF,
    ".docx": DocumentFormat.DOCX,
    ".md": DocumentFormat.MARKDOWN,
    ".markdown": DocumentFormat.MARKDOWN,
    ".html": DocumentFormat.HTML,
    ".htm": DocumentFormat.HTML,
    ".txt": DocumentFormat.TEXT,
}


class LoaderFactory:
    _loaders: ClassVar[dict[DocumentFormat, type[DocumentLoader]]] = {
        DocumentFormat.PDF: PdfLoader,
        DocumentFormat.DOCX: DocxLoader,
        DocumentFormat.HTML: HtmlLoader,
        DocumentFormat.MARKDOWN: MarkdownLoader,
        DocumentFormat.TEXT: TextLoader,
    }

    @classmethod
    def for_format(cls, format: DocumentFormat) -> DocumentLoader:
        loader_cls = cls._loaders.get(format)
        if loader_cls is None:
            raise LoaderError(f"unsupported document format: {format.value}")
        return loader_cls()

    @staticmethod
    def detect_format(filename: str) -> DocumentFormat:
        name = PureWindowsPath(filename).name
        suffix = PurePosixPath(name).suffix.lower()
        format = _SUFFIX_FORMATS.get(suffix)
        if format is None:
            raise LoaderError(f"unsupported file extension: {filename}")
        return format
